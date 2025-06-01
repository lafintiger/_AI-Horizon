#!/usr/bin/env python3
"""
Document and Media Processors for Manual Entry

Handles extraction of text content from various file formats and media sources.
"""

import os
import re
import json
from pathlib import Path
from typing import Optional, Dict, Any
import tempfile

from aih.utils.logging import get_logger

logger = get_logger(__name__)

def process_document(file_path: Path) -> str:
    """
    Process a document file and extract text content.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text content
    """
    try:
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.txt':
            return process_txt_file(file_path)
        elif file_extension == '.pdf':
            return process_pdf_file(file_path)
        elif file_extension in ['.docx', '.doc']:
            return process_docx_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        logger.error(f"Error processing document {file_path}: {e}")
        return f"Error processing document: {str(e)}"

def process_txt_file(file_path: Path) -> str:
    """Extract text from a TXT file."""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.info(f"Successfully read TXT file with {encoding} encoding")
                return content.strip()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode file with any supported encoding")
        
    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {e}")
        raise

def process_pdf_file(file_path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        # Try PyPDF2 first (lighter dependency)
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_content = []
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())
                
                content = '\n'.join(text_content).strip()
                if content:
                    logger.info(f"Successfully extracted text from PDF using PyPDF2")
                    return content
        
        except ImportError:
            logger.warning("PyPDF2 not available, trying pdfplumber")
        
        # Fallback to pdfplumber (better text extraction)
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                text_content = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                
                content = '\n'.join(text_content).strip()
                logger.info(f"Successfully extracted text from PDF using pdfplumber")
                return content
                
        except ImportError:
            logger.error("Neither PyPDF2 nor pdfplumber available for PDF processing")
            raise ImportError("PDF processing requires PyPDF2 or pdfplumber. Install with: pip install PyPDF2 pdfplumber")
        
    except Exception as e:
        logger.error(f"Error processing PDF file {file_path}: {e}")
        raise

def process_docx_file(file_path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        import docx
        
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        content = '\n'.join(text_content).strip()
        logger.info(f"Successfully extracted text from DOCX file")
        return content
        
    except ImportError:
        logger.error("python-docx not available for DOCX processing")
        raise ImportError("DOCX processing requires python-docx. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"Error processing DOCX file {file_path}: {e}")
        raise

def extract_youtube_transcript(url: str) -> Optional[Dict[str, Any]]:
    """
    Extract transcript from a YouTube video.
    
    Args:
        url: YouTube video URL
        
    Returns:
        Dictionary with transcript data or None if extraction fails
    """
    try:
        # Try youtube-transcript-api first (most reliable)
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            import re
            
            # Extract video ID from URL
            video_id = extract_video_id_from_url(url)
            if not video_id:
                raise ValueError("Could not extract video ID from URL")
            
            # Get transcript
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
            
            # Combine transcript segments
            full_transcript = ' '.join([item['text'] for item in transcript_list])
            
            # Try to get video metadata
            video_data = get_youtube_metadata(video_id)
            
            result = {
                'transcript': full_transcript,
                'video_id': video_id,
                'title': video_data.get('title', 'Unknown Title'),
                'channel': video_data.get('channel', 'Unknown Channel'),
                'duration': video_data.get('duration', 'Unknown Duration'),
                'url': url
            }
            
            logger.info(f"Successfully extracted YouTube transcript for video {video_id}")
            return result
            
        except ImportError:
            logger.warning("youtube-transcript-api not available, trying yt-dlp")
        
        # Fallback to yt-dlp
        try:
            import yt_dlp
            
            ydl_opts = {
                'quiet': True,
                'no_warnings': True,
                'writeautomaticsub': True,
                'writesubtitles': True,
                'skip_download': True,
                'subtitleslangs': ['en', 'en-US', 'en-GB'],
            }
            
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)
                
                # Get basic metadata
                result = {
                    'video_id': info.get('id', ''),
                    'title': info.get('title', 'Unknown Title'),
                    'channel': info.get('uploader', 'Unknown Channel'),
                    'duration': str(info.get('duration', 'Unknown Duration')),
                    'url': url,
                    'transcript': ''
                }
                
                # Try to get subtitles/transcript
                subtitles = info.get('subtitles', {}) or info.get('automatic_captions', {})
                
                for lang in ['en', 'en-US', 'en-GB']:
                    if lang in subtitles:
                        # This is a simplified approach - would need more complex processing
                        # for actual subtitle file parsing
                        result['transcript'] = "Transcript extraction via yt-dlp requires additional processing"
                        break
                
                logger.info(f"Extracted YouTube metadata using yt-dlp")
                return result
                
        except ImportError:
            logger.error("Neither youtube-transcript-api nor yt-dlp available")
            raise ImportError("YouTube processing requires youtube-transcript-api or yt-dlp. Install with: pip install youtube-transcript-api yt-dlp")
        
    except Exception as e:
        logger.error(f"Error extracting YouTube transcript from {url}: {e}")
        return None

def extract_video_id_from_url(url: str) -> Optional[str]:
    """Extract YouTube video ID from various URL formats."""
    patterns = [
        r'(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/)([a-zA-Z0-9_-]{11})',
        r'youtube\.com/.*[?&]v=([a-zA-Z0-9_-]{11})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    
    return None

def get_youtube_metadata(video_id: str) -> Dict[str, str]:
    """
    Get basic metadata for a YouTube video.
    
    Args:
        video_id: YouTube video ID
        
    Returns:
        Dictionary with metadata
    """
    # This is a placeholder - in a real implementation, you might use:
    # - YouTube Data API v3
    # - yt-dlp for metadata extraction
    # - Web scraping (not recommended)
    
    return {
        'title': f'Video {video_id}',
        'channel': 'Unknown Channel',
        'duration': 'Unknown Duration'
    }

def install_missing_dependencies():
    """
    Check for and install missing dependencies for document processing.
    """
    missing_deps = []
    
    # Check for PDF processing
    try:
        import PyPDF2
    except ImportError:
        try:
            import pdfplumber
        except ImportError:
            missing_deps.append("PyPDF2 or pdfplumber (for PDF processing)")
    
    # Check for DOCX processing
    try:
        import docx
    except ImportError:
        missing_deps.append("python-docx (for DOCX processing)")
    
    # Check for YouTube processing
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
    except ImportError:
        try:
            import yt_dlp
        except ImportError:
            missing_deps.append("youtube-transcript-api or yt-dlp (for YouTube processing)")
    
    if missing_deps:
        print("⚠️  Missing dependencies for manual entry:")
        for dep in missing_deps:
            print(f"   • {dep}")
        print("\n📦 Install missing dependencies:")
        print("pip install PyPDF2 pdfplumber python-docx youtube-transcript-api yt-dlp")
        return False
    
    return True

if __name__ == "__main__":
    # Test dependency availability
    print("🔍 Checking manual entry dependencies...")
    if install_missing_dependencies():
        print("✅ All dependencies available!")
    else:
        print("❌ Some dependencies missing. Please install them for full functionality.") 